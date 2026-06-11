# -*- coding: utf-8 -*-
"""Generate updated design and requirements Word documents."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_design_doc():
    doc = Document()
    set_doc_font(doc)

    add_title(doc, '"基于多情景房源查询及可视化数据网页"设计方案')
    add_body(doc, "（版本号：V2.0）")
    doc.add_paragraph()

    add_h1(doc, "文件修订记录")
    table = doc.add_table(rows=4, cols=5)
    hdr = ["版本编号", "变化状态", "简要说明", "日期", "变更人"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ["V1.0", "N", "新建（初始网页搭建）", "2026/3/22", "曾嘉明、梁东海"],
        ["V1.1", "C", "交互功能实现与全部功能网页初步完善", "2026/4/23", "曾嘉明、梁东海"],
        ["V2.0", "C", "房价预测升级为多元线性回归（OLS）；散点图与拟合线完善；咨询栏优化；文档与代码对齐", "2026/6/11", "曾嘉明、梁东海"],
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    add_body(doc, "变化状态：N—新建，C—变更，D—删除")
    doc.add_paragraph()

    add_h1(doc, "1 简介")
    add_h2(doc, "1.1 目的")
    add_body(doc, '本项目为基于 JSP/Servlet 的粤港澳大湾区房产信息分析系统，运行环境为 Eclipse + Apache Tomcat 9.0，数据库采用 XAMPP 提供的 MySQL（端口 3307，库名 housing_db）。系统围绕「房源搜索、价格走势图、房价预测」三大核心功能构建，实现从多条件筛选、趋势可视化到基于数据库训练的租售分模型房价预测的完整业务流程。')
    add_body(doc, "当前版本已完成：① 租房/购房双模式数据隔离；② Canvas 价格走势折线图；③ 多元线性回归（OLS）房价预测，含 R²、边际系数、四宫格散点图与拟合线、数据驱动建议文案；④ 浏览器本地收藏。")

    add_h2(doc, "1.2 适用范围")
    add_body(doc, "项目组内部使用、课程答辩与演示。")

    add_h2(doc, "1.3 术语表")
    add_bullets(doc, [
        "listing_type：房源类型，rent（租房）/ sale（购房）",
        "OLS：最小二乘线性回归，用于拟合房价预测方程",
        "R²：决定系数，衡量回归模型拟合优度",
        "sessionStorage：浏览器会话存储，保存租购模式 listingType",
        "localStorage：浏览器本地存储，保存用户收藏房源",
        "Canvas：HTML5 画布，用于走势图与预测散点图绘制",
    ])

    add_h2(doc, "1.4 引用文件")
    add_body(doc, "《基于多情景房源查询及数据可视化网页需求说明书》V2.0")

    add_h1(doc, "2 关键功能流程")
    add_h2(doc, "2.1 总体流程")
    add_body(doc, "用户进入首页 → 选择【租房】或【购房】→ 选择身份板块（租客/买方 或 房东/售房方，仅用于引导）→ 进入房源搜索 / 价格走势图 / 房价预测。租购模式写入 sessionStorage，后续各功能接口均按 listing_type 过滤数据。")

    add_h2(doc, "2.2 房源搜索")
    add_body(doc, "用户在 search.jsp 填写最多 8 项可选条件（地区、年份、月份、预算、面积、楼层、距地铁、房龄），POST /api/search 提交。后端动态拼接 SQL（仅对非空条件加 WHERE），参数化防注入，按年月降序、价格升序，最多返回 120 条。结果以 JSON 渲染卡片，点击 + 写入 localStorage 收藏。")

    add_h2(doc, "2.3 价格走势图")
    add_body(doc, "用户在 trend.jsp 选择地区（必选）和年份（可选），GET /api/trend。仅选地区时按年份聚合 AVG(price/area)；选地区+年份时按 1~12 月聚合。前端 Canvas 绘制坐标轴、折线动画与数据点。")

    add_h2(doc, "2.4 房价预测（V2.0 线性回归）")
    add_body(doc, "用户在 predict.jsp 填写 7 项必填特征（地区、年份、月份、面积、楼层、距地铁、房龄），POST /api/predict。后端流程：")
    add_bullets(doc, [
        "首次预测时，RegressionService 从数据库分别加载全部 rent / sale 记录训练两套 OLS 模型并缓存在内存；",
        "特征向量：截距 + 年份、月份、面积、楼层、房龄、距地铁 + 8 个城市哑变量（广州为基准）；",
        "预测目标为总价 price（万），展示单价 = 预测总价 / 面积；",
        "返回 R²、训练样本数、边际系数 β、四组拟合线数据、200 条散点样本、建议文案；",
        "前端 2×2 Canvas 展示面积/楼层/距地铁/房龄与价格散点及蓝色拟合线，支持分步动画。",
    ])
    add_body(doc, "建议文案由 SuggestionService 基于同地区同月市场均价及同比变化生成，与回归预测相互独立。")

    add_h2(doc, "2.3 关键功能点与说明")
    add_bullets(doc, [
        "租售分模型：同一特征结构，训练数据按 listing_type 完全分离；",
        "全量训练：使用库内该类型全部有效记录拟合，而非局部小样本；",
        "建议服务：按地区+年月+面积±15% 查询市场统计，样本不足时降级为地区范围。",
    ])

    add_h1(doc, "3 架构设计")
    add_h2(doc, "3.1 系统总体结构（四层）")
    add_bullets(doc, [
        "表示层：index.jsp、search.jsp、trend.jsp、predict.jsp、favorites.jsp；assets/css/style.css、assets/js/app.js",
        "控制层：SearchServlet、TrendServlet、PredictServlet、ScatterServlet；WEB-INF/web.xml",
        "业务层：RegressionService、PredictionService、SuggestionService",
        "数据访问层：HouseListingDAO、DBUtil；模型层：HouseListing、LinearRegressionModel、MarketStats；工具：ListingTypeUtil",
        "数据层：MySQL housing_db，单表 house_listings",
    ])
    add_body(doc, "链路：JSP/JS → Servlet → Service → DAO → MySQL → JSON 返回 → 页面渲染。")

    add_h2(doc, "3.2 核心类职责")
    add_bullets(doc, [
        "LinearRegressionModel：高斯消元求解 OLS 系数，提供 predict、R²",
        "RegressionService：特征工程、双模型训练缓存、拟合线、边际系数",
        "PredictionService：预测门面，调用 RegressionService",
        "SuggestionService：市场均价对比与同比建议规则",
        "HouseListingDAO：搜索、趋势聚合、训练数据、散点采样、市场统计",
    ])

    add_h1(doc, "4 接口设计")
    add_h2(doc, "4.1 POST /api/search")
    add_body(doc, "参数（均可选组合）：district, year, month, budget, area, floor, distanceToSubway, houseAge, listingType。返回 success、data[]（district, year, month, price, area, floor）。")

    add_h2(doc, "4.2 GET /api/trend")
    add_body(doc, "参数：district（必填）、listingType（必填）、year（可选）。返回 success、mode（year/month）、points。")

    add_h2(doc, "4.3 POST /api/predict")
    add_body(doc, "参数（必填）：district, year, month, area, floor, distanceToSubway, houseAge, listingType。返回 unitPrice, totalPrice, unitLabel, rSquared, trainSampleCount, modelType, fitLines, scatterPoints, scatterCount, impact[], suggestion。")

    add_h2(doc, "4.4 GET /api/scatter")
    add_body(doc, "参数：listingType。返回最多 200 条随机散点（area, floor, distance, age, price），供预测页图表回退加载。")

    add_h1(doc, "5 界面设计")
    add_bullets(doc, [
        "首页 index.jsp：租购模式、身份板块、三功能入口、底部多资讯滚动栏（JS 驱动从右向左）",
        "房源搜索 search.jsp：左 8 项筛选 + 右结果卡片 + 收藏",
        "价格走势图 trend.jsp：左地区/年份 + 右 Canvas 折线图",
        "房价预测 predict.jsp：左 7 项输入 + 右四宫格散点/拟合线、边际系数表、建议区",
        "用户收藏 favorites.jsp：localStorage 列表与删除",
        "无登录页；整体蓝白简洁风格，顶部导航可跨页跳转",
    ])

    add_h1(doc, "6 数据结构设计")
    add_h2(doc, "6.1 数据库")
    add_body(doc, "数据库名：housing_db。核心表：house_listings（单表方案，无用户表、收藏表、预测日志表）。")
    add_bullets(doc, [
        "id BIGINT PK AUTO_INCREMENT",
        "district VARCHAR(20) 地区（广州、深圳、珠海、佛山、东莞、中山、惠州、江门、肇庆）",
        "year INT、month INT（1~12）",
        "price DECIMAL(10,2) 总价（万）；租房为月租总额，购房为售价",
        "area DECIMAL(8,2) 面积（㎡）",
        "floor INT 楼层",
        "house_age DECIMAL(6,2) 房龄（年）",
        "distance_to_subway DECIMAL(6,2) 距地铁（km）",
        "listing_type VARCHAR(10) rent / sale",
        "索引 idx_qry(district, year, month, listing_type)",
    ])
    add_body(doc, "数据量参考：租售各 1.6 万+ 条可满足回归训练（特征约 15 维，样本量远大于参数数）。")

    add_h2(doc, "6.2 前后端状态")
    add_bullets(doc, [
        "sessionStorage.listingType：rent 或 sale",
        "localStorage.favorites：收藏房源 JSON 数组",
        "回归模型：JVM 内存缓存，Tomcat 重启后重新训练",
    ])

    add_h2(doc, "6.3 配置文件")
    add_bullets(doc, [
        "WEB-INF/web.xml：Servlet 映射",
        "DBUtil.java：jdbc:mysql://localhost:3307/housing_db，用户 root",
        "sql/init.sql：建库建表脚本",
    ])

    out = r"d:\SoftwareClass\Project\.docs\基于多情景房源查询与可视化数据网页设计说明_V2.0.docx"
    doc.save(out)
    return out


def build_requirements_doc():
    doc = Document()
    set_doc_font(doc)

    add_title(doc, '"基于多情景房源查询及可视化数据网页"需求说明书')
    add_body(doc, "（版本号：V2.0）")
    doc.add_paragraph()

    add_h1(doc, "文件修订记录")
    table = doc.add_table(rows=3, cols=5)
    hdr = ["版本编号", "变化状态", "简要说明", "日期", "变更人"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ["V1.0", "N", "新建及三大功能基础搭建", "2026-03-22", "曾嘉明、梁东海"],
        ["V2.0", "C", "对齐最终实现：线性回归预测、散点拟合线、咨询栏、删除无效章节", "2026-06-11", "曾嘉明、梁东海"],
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()

    add_h1(doc, "1 简介")
    add_h2(doc, "1.1 文档介绍")
    add_body(doc, "本文档描述智慧房源平台（房产信息服务平台）的功能需求、操作流程与非功能要求，供设计、实现与验收使用。")
    add_h2(doc, "1.2 开发背景")
    add_body(doc, "面向粤港澳大湾区租房与购房人群，提供房源多条件检索、均价走势查看及基于历史数据的线性回归房价预测，辅助用户了解市场并辅助决策。")

    add_h1(doc, "2 网页功能介绍")
    add_h2(doc, "2.1 操作总览")
    add_body(doc, "首页选择租购模式与身份 → 进入三大功能之一 → 各功能页顶栏可互相跳转、重试或回首页。使用前必须在首页选择【租房】或【购房】，否则预测等功能将提示返回首页。")

    add_h2(doc, "2.2 首页")
    add_bullets(doc, [
        "【用户收藏】：跳转 favorites.jsp，查看/删除 localStorage 收藏",
        "【租房】【购房】：设置 sessionStorage.listingType",
        "身份板块：租客/买方、房东/售房方（界面引导，不参与后端业务分支）",
        "功能入口：房源搜索、价格走势图、房价预测",
        "底部资讯栏：8 条随机楼市资讯，自屏幕右侧向左循环滚动",
    ])

    add_h2(doc, "2.3 功能一：房源搜索")
    add_body(doc, "8 项可选特征：地区、年份、月份、预算（万）、面积（㎡）、楼层、距地铁（km）、房龄（年）。地区为珠三角 9 市（广州、深圳、珠海、佛山、东莞、中山、惠州、江门、肇庆）。所有数据存于同一数据库同一表，通过 listing_type 区分租售。")
    add_body(doc, "点击【确定】后右侧显示匹配房源（最多 120 条），格式示例：广州 | 2025年8月 | 86㎡ | 158万。点击 + 加入收藏。")

    add_h2(doc, "2.4 功能二：价格走势图")
    add_body(doc, "左侧选择地区（必选）、年份（可选）。点击【查看】后右侧 Canvas 显示均价走势：")
    add_bullets(doc, [
        "仅选地区：横轴为年份，纵轴为 AVG(price/area)（万/㎡或万/㎡/月）；",
        "选地区+年份：横轴为 1~12 月；",
        "折线与数据点带动画绘制。",
    ])

    add_h2(doc, "2.5 功能三：房价预测")
    add_body(doc, "7 项必填特征：地区、年份、月份、面积、楼层、距地铁、房龄。地区下拉为 8 市（肇庆未列入下拉，后端仍支持）。点击【确定】后在同一页面展示：")
    add_bullets(doc, [
        "模型信息：多元线性回归(OLS)、R²、训练样本数、散点样本数；",
        "四宫格散点图：面积/楼层/距地铁/房龄 vs 价格，红色散点 + 蓝色拟合线；",
        "边际系数表：各特征每增加 1 单位的总价变化（万）；",
        "预测月租或预测总价、单位价格；",
        "建议文本：与数据库同地区同月均价、同比变化对比（2023 年及以前无同比建议）。",
    ])
    add_body(doc, "预测算法需求：必须使用数据库中 rent/sale 各自全量样本训练独立线性回归方程，不得仅用手写规则或局部均值代替。")

    add_h1(doc, "3 系统通用性需求")
    add_h2(doc, "3.1 性能需求")
    add_body(doc, "常规搜索、走势查询响应时间 &lt; 3 秒。首次房价预测因需训练模型（约 1.6 万+ 条）允许 5~15 秒，之后同会话内预测应明显加快（模型已缓存）。")
    add_h2(doc, "3.2 安全需求")
    add_body(doc, "课程演示级：无用户登录；SQL 使用 PreparedStatement；数据库账号配置在服务端。不承诺等保三级。")
    add_h2(doc, "3.3 数据需求")
    add_body(doc, "房源数据存储于本地 MySQL；需保证 rent、sale 两类均有足够样本（建议每类不少于数百条，实际约 1.6 万+）以支撑回归训练。")
    add_h2(doc, "3.4 操作性需求")
    add_bullets(doc, [
        "Windows 10/11 浏览器访问（Chrome、Edge 等）",
        "XAMPP MySQL（3307）+ Tomcat 9 部署",
        "界面中文，UTF-8 编码",
        "各页提供返回首页、重试及功能互跳",
    ])
    add_h2(doc, "3.5 兼容性")
    add_body(doc, "桌面浏览器宽屏与全屏下，咨询栏应铺满视口宽度并自右向左滚动。")

    add_h1(doc, "4 系统功能用例")
    add_h2(doc, "4.1 角色说明")
    add_bullets(doc, [
        "访客用户：使用搜索、走势、预测、收藏（无账号体系）",
        "系统：训练回归模型、聚合统计、生成 JSON 与建议",
    ])

    add_h2(doc, "4.2 用例列表")
    cases = [
        ("UC-01 选择租购模式", "访客", "进入首页", "sessionStorage 写入 listingType", "未选择则功能页提示返回首页"),
        ("UC-02 条件搜索房源", "访客", "已选租购模式", "展示 ≤120 条房源卡片", "无结果时显示空列表"),
        ("UC-03 收藏房源", "访客", "搜索有结果", "写入 localStorage", "收藏页可删除"),
        ("UC-04 查看价格走势", "访客", "已选租购模式、已选地区", "Canvas 折线图", "无数据时提示"),
        ("UC-05 线性回归预测", "访客", "已选租购模式、填齐 7 项", "返回价格、R²、系数、散点与拟合线、建议", "数据库不可用或样本不足时失败提示"),
        ("UC-06 查看收藏列表", "访客", "—", "favorites.jsp 列表", "无收藏时为空"),
    ]
    t = doc.add_table(rows=1, cols=5)
    headers = ["用例编号", "名称", "执行者", "前置条件", "结果/异常"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for case in cases:
        row = t.add_row().cells
        for i, v in enumerate(case):
            row[i].text = v

    add_h2(doc, "4.3 房价预测用例补充说明")
    add_bullets(doc, [
        "输入：7 项特征 + 当前租购模式",
        "处理：训练/加载对应 OLS 模型 → 预测总价 → 生成拟合线与散点 → 查询市场统计写建议",
        "输出：unitPrice、totalPrice、impact[]、fitLines、scatterPoints、suggestion",
        "验收标准：租售切换后预测结果不同；R² 与训练样本数有合理数值；散点图可见多点与拟合线",
    ])

    out = r"d:\SoftwareClass\Project\.docs\基于多情景房源查询与数据可视化网页需求说明书_V2.0.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    d1 = build_design_doc()
    d2 = build_requirements_doc()
    print("Generated:", d1)
    print("Generated:", d2)
